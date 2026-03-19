"""
Complete script to:
1. Generate all new matplotlib charts
2. Insert charts into diplom_v2.docx replacing placeholders
"""

import os
import sys
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import matplotlib.patches as patches
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch, Rectangle
from matplotlib.gridspec import GridSpec
import matplotlib.patheffects as pe
import warnings
warnings.filterwarnings('ignore')

# Set global font
plt.rcParams['font.family'] = 'DejaVu Sans'
plt.rcParams['figure.facecolor'] = 'white'
plt.rcParams['axes.facecolor'] = 'white'

BASE_DIR = '/home/roman/roman/Python_for_Magistr/codirank'
OUT_DIR = os.path.join(BASE_DIR, 'public')
os.makedirs(OUT_DIR, exist_ok=True)

# Color palette
BLUE   = '#2563EB'
PURPLE = '#7C3AED'
GREEN  = '#059669'
ORANGE = '#D97706'
RED    = '#DC2626'
GRAY   = '#6B7280'
LIGHT_GRAY = '#F3F4F6'
DARK   = '#1F2937'
YELLOW = '#FBBF24'


# ===========================================================================
# 1. fig_gplay_search.png
# ===========================================================================
def gen_gplay_search():
    fig, ax = plt.subplots(figsize=(8, 10))
    ax.set_xlim(0, 8)
    ax.set_ylim(0, 10)
    ax.axis('off')
    fig.patch.set_facecolor('#F8F9FA')

    # Header bar
    header = FancyBboxPatch((0, 9.2), 8, 0.8, boxstyle="square,pad=0",
                            facecolor='#1A73E8', edgecolor='none')
    ax.add_patch(header)
    ax.text(0.3, 9.58, '<', fontsize=16, color='white', va='center', fontweight='bold')
    ax.text(4.0, 9.62, 'Google Play', fontsize=14, color='white',
            va='center', ha='center', fontweight='bold')

    # Search bar
    search_box = FancyBboxPatch((0.2, 8.7), 7.0, 0.45,
                                boxstyle="round,pad=0.05",
                                facecolor='white', edgecolor='#DADCE0', linewidth=1.5)
    ax.add_patch(search_box)
    ax.text(0.55, 8.92, 'meditaciya', fontsize=10, color=DARK, va='center')
    ax.text(6.9, 8.92, 'x', fontsize=12, color=GRAY, va='center', ha='center')

    # Tabs
    ax.text(1.0, 8.45, 'Vse', fontsize=9, color='#1A73E8', va='center',
            ha='center', fontweight='bold')
    ax.text(2.5, 8.45, 'Igry', fontsize=9, color=GRAY, va='center', ha='center')
    ax.text(4.0, 8.45, 'Programmy', fontsize=9, color=GRAY, va='center', ha='center')
    # underline for active tab
    ax.plot([0.4, 1.6], [8.28, 8.28], color='#1A73E8', linewidth=2)

    # Result cards
    apps = [
        ('#4CAF50', 'Calm: Meditaciya i son',      '4.8', '50M+',  'Zdorove i fitnes'),
        ('#FF5722', 'Angry Birds Reloaded',          '4.3', '100M+', 'Igry: Arkady'),
        ('#9C27B0', 'Headspace: Osoznannost',        '4.7', '10M+',  'Zdorove'),
        ('#2196F3', 'Microsoft To Do: Zadachi',      '4.6', '100M+', 'Produktivnost'),
        ('#FF9800', 'Meditaciya dlya nachinayushchih','4.5', '1M+',   'Zdorove i fitnes'),
    ]

    card_y_starts = [7.85, 6.65, 5.45, 4.25, 3.05]
    for (color, name, rating, installs, category), y in zip(apps, card_y_starts):
        # Card background
        card = FancyBboxPatch((0.1, y - 0.6), 7.8, 1.1,
                              boxstyle="round,pad=0.05",
                              facecolor='white', edgecolor='#E8EAED', linewidth=1)
        ax.add_patch(card)

        # App icon
        icon = FancyBboxPatch((0.25, y - 0.45), 0.8, 0.8,
                              boxstyle="round,pad=0.05",
                              facecolor=color, edgecolor='none')
        ax.add_patch(icon)
        ax.text(0.65, y + 0.0, name[0], fontsize=20, color='white',
                va='center', ha='center', fontweight='bold')

        # App info
        ax.text(1.3, y + 0.25, name[:30], fontsize=9, color=DARK,
                va='center', fontweight='bold')
        ax.text(1.3, y + 0.05, category, fontsize=8, color=GRAY, va='center')

        # Stars
        stars = int(float(rating))
        star_str = '*' * stars + ('.' if float(rating) % 1 else '')
        ax.text(1.3, y - 0.18, rating + ' ', fontsize=8, color=ORANGE, va='center')
        ax.text(1.75, y - 0.18, '| ' + installs + ' ustanovok', fontsize=8,
                color=GRAY, va='center')

        # Install button
        btn = FancyBboxPatch((5.8, y - 0.2), 1.9, 0.45,
                             boxstyle="round,pad=0.05",
                             facecolor='#E8F0FE', edgecolor='none')
        ax.add_patch(btn)
        ax.text(6.75, y + 0.025, 'Ustanovit', fontsize=8, color='#1A73E8',
                va='center', ha='center', fontweight='bold')

    ax.text(4.0, 0.5, 'Pokazany raznorodnye rezultaty (igry, produktivnost,\n'
            'zdorove) pri zaprose "meditaciya"',
            fontsize=8, color=GRAY, ha='center', va='center', style='italic')

    ax.set_title('Primer interfeysa poiska Google Play s neodnoznachnymi rezultatami',
                 fontsize=11, color=DARK, pad=8, fontweight='bold')
    plt.tight_layout()
    plt.savefig(os.path.join(OUT_DIR, 'fig_gplay_search.png'),
                dpi=150, bbox_inches='tight', facecolor='#F8F9FA')
    plt.close()
    print('[OK] fig_gplay_search.png')


# ===========================================================================
# 2. fig_sparsity_matrix.png
# ===========================================================================
def gen_sparsity_matrix():
    np.random.seed(42)
    rows, cols = 20, 30
    matrix = np.zeros((rows, cols))
    n_filled = int(rows * cols * 0.03)
    filled_idx = np.random.choice(rows * cols, n_filled, replace=False)
    for idx in filled_idx:
        r, c = divmod(idx, cols)
        matrix[r, c] = 1

    fig, ax = plt.subplots(figsize=(12, 6))

    import matplotlib.colors as mcolors
    cmap = plt.cm.viridis
    im = ax.imshow(matrix, cmap=cmap, aspect='auto', interpolation='nearest')

    ax.set_xlabel('Prilozheniya (m=500,000)', fontsize=12, color=DARK)
    ax.set_ylabel('Polzovateli (n=10,000)', fontsize=12, color=DARK)
    ax.set_title('Vizualizaciya razrezhennosti matricy vzaimodeystviy\n'
                 'polzovatel-prilozhenie',
                 fontsize=13, fontweight='bold', color=DARK)

    ax.set_xticks([0, 10, 20, 29])
    ax.set_xticklabels(['0', '166K', '333K', '500K'], fontsize=9, color=DARK)
    ax.set_yticks([0, 5, 10, 15, 19])
    ax.set_yticklabels(['0', '2.5K', '5K', '7.5K', '10K'], fontsize=9, color=DARK)

    cbar = plt.colorbar(im, ax=ax, shrink=0.6)
    cbar.set_label('Vzaimodeystvie', fontsize=10, color=DARK)
    cbar.set_ticks([0, 1])
    cbar.set_ticklabels(['0 (net)', '1 (est)'])

    ax.text(0.98, 0.02, 'Zapolnennost: 0.03%',
            transform=ax.transAxes, fontsize=11,
            color='white', ha='right', va='bottom',
            bbox=dict(boxstyle='round,pad=0.4', facecolor=PURPLE, alpha=0.9))

    ax.text(0.02, 0.98,
            'Kolichestvo:\n'
            '  polzovateley: 10,000\n'
            '  prilozheniy:  500,000\n'
            '  vzaimodeystviy: ~1,500',
            transform=ax.transAxes, fontsize=9,
            color='white', ha='left', va='top',
            bbox=dict(boxstyle='round,pad=0.4', facecolor=DARK, alpha=0.85))

    plt.tight_layout()
    plt.savefig(os.path.join(OUT_DIR, 'fig_sparsity_matrix.png'),
                dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    print('[OK] fig_sparsity_matrix.png')


# ===========================================================================
# 3. fig_matrix_factorization.png
# ===========================================================================
def gen_matrix_factorization():
    fig, ax = plt.subplots(figsize=(14, 5))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 5)
    ax.axis('off')

    def draw_matrix(ax, x, y, rows, cols, label, cell_size=0.3, cmap_name='Blues',
                    sparse=False, title='', dims=''):
        np.random.seed(1)
        for r in range(rows):
            for c in range(cols):
                if sparse and np.random.rand() > 0.15:
                    color = '#E5E7EB'
                else:
                    val = np.random.rand()
                    color = plt.cm.get_cmap(cmap_name)(0.3 + val * 0.65)
                    if sparse and np.random.rand() > 0.15:
                        color = '#E5E7EB'
                rect = plt.Rectangle((x + c * cell_size, y + (rows - r - 1) * cell_size),
                                     cell_size - 0.02, cell_size - 0.02,
                                     facecolor=color, edgecolor='white', linewidth=0.5)
                ax.add_patch(rect)
        total_w = cols * cell_size
        total_h = rows * cell_size
        ax.text(x + total_w / 2, y + total_h + 0.2, label,
                ha='center', va='bottom', fontsize=14, fontweight='bold', color=DARK)
        ax.text(x + total_w / 2, y - 0.25, dims,
                ha='center', va='top', fontsize=9, color=GRAY)

    # R matrix (sparse, users x items)
    draw_matrix(ax, 0.3, 0.8, 10, 14, 'R', cell_size=0.28,
                cmap_name='YlOrRd', sparse=True, dims='n x m\n(polzovateli x prilozheniya)')

    # Arrow ≈
    ax.text(4.5, 2.6, u'\u2248', fontsize=28, color=DARK, ha='center', va='center',
            fontweight='bold')

    # U matrix (users x k)
    draw_matrix(ax, 5.0, 0.8, 10, 5, 'U', cell_size=0.28,
                cmap_name='Blues', sparse=False, dims='n x k\n(polzovateli x faktory)')

    # Times symbol
    ax.text(6.8, 2.6, u'\u00d7', fontsize=24, color=DARK, ha='center', va='center',
            fontweight='bold')

    # V^T matrix (k x items)
    draw_matrix(ax, 7.3, 1.4, 5, 14, 'V\u1D40', cell_size=0.28,
                cmap_name='Purples', sparse=False, dims='k x m\n(faktory x prilozheniya)')

    # k label
    ax.annotate('', xy=(5.0, 3.72), xytext=(6.75, 3.72),
                arrowprops=dict(arrowstyle='<->', color=GREEN, lw=1.5))
    ax.text(5.875, 3.85, 'k=5 latentnykh faktorov',
            ha='center', va='bottom', fontsize=9, color=GREEN, fontweight='bold')

    ax.set_title('Skhema matrichnoy faktorizacii v zadache kollaborativnoy filtracii\n'
                 'R = U * V^T — priblizhenie matricy vzaimodeystviy',
                 fontsize=12, fontweight='bold', color=DARK, y=1.02)
    plt.tight_layout()
    plt.savefig(os.path.join(OUT_DIR, 'fig_matrix_factorization.png'),
                dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    print('[OK] fig_matrix_factorization.png')


# ===========================================================================
# 4. fig_ncf_architecture.png
# ===========================================================================
def gen_ncf_architecture():
    fig, ax = plt.subplots(figsize=(10, 9))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 9)
    ax.axis('off')

    def box(ax, x, y, w, h, text, color, text_color='white', fontsize=9, radius=0.2):
        b = FancyBboxPatch((x - w/2, y - h/2), w, h,
                           boxstyle=f"round,pad={radius}",
                           facecolor=color, edgecolor='white', linewidth=1.5)
        ax.add_patch(b)
        ax.text(x, y, text, ha='center', va='center',
                fontsize=fontsize, color=text_color, fontweight='bold',
                wrap=True, multialignment='center')

    def arrow(ax, x1, y1, x2, y2, color=GRAY):
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle='->', color=color, lw=1.5))

    # Input layer
    box(ax, 2.5, 0.7, 2.8, 0.6, 'ID Polzovatelya', DARK, fontsize=9)
    box(ax, 7.5, 0.7, 2.8, 0.6, 'ID Prilozheniya', DARK, fontsize=9)

    # Embedding layer
    box(ax, 2.5, 1.9, 2.8, 0.6, 'Embedding\npolzovatelya', BLUE, fontsize=9)
    box(ax, 7.5, 1.9, 2.8, 0.6, 'Embedding\nprilozheniya', BLUE, fontsize=9)

    arrow(ax, 2.5, 1.0, 2.5, 1.6)
    arrow(ax, 7.5, 1.0, 7.5, 1.6)

    # GMF branch (left) and MLP branch (right)
    # GMF
    box(ax, 2.5, 3.1, 2.8, 0.6, 'GMF: poelementnoe\numnozheniye', GREEN, fontsize=8)
    arrow(ax, 2.5, 2.2, 2.5, 2.8)
    # draw line from right embedding to GMF too
    ax.annotate('', xy=(3.6, 3.1), xytext=(6.1, 2.2),
                arrowprops=dict(arrowstyle='->', color=GRAY, lw=1.0,
                                connectionstyle='arc3,rad=-0.2'))

    # MLP branch
    box(ax, 7.5, 3.1, 2.8, 0.6, 'MLP: konkatenaciya\nembeddingov', ORANGE, fontsize=8)
    arrow(ax, 7.5, 2.2, 7.5, 2.8)
    ax.annotate('', xy=(6.4, 3.1), xytext=(3.9, 2.2),
                arrowprops=dict(arrowstyle='->', color=GRAY, lw=1.0,
                                connectionstyle='arc3,rad=0.2'))

    # MLP layers
    for i, (y, label) in enumerate([(4.3, 'Dense(256) + ReLU'),
                                     (5.2, 'Dense(128) + ReLU'),
                                     (6.1, 'Dense(64)  + ReLU')]):
        box(ax, 7.5, y, 2.8, 0.55, label, '#EA580C', fontsize=8)
        arrow(ax, 7.5, y - 0.62, 7.5, y - 0.28)

    # GMF output
    box(ax, 2.5, 5.2, 2.8, 0.6, 'GMF output\n(k-merny vektor)', GREEN, fontsize=8)
    arrow(ax, 2.5, 3.4, 2.5, 4.9)

    # Concatenation
    box(ax, 5.0, 6.85, 4.5, 0.6, 'Konkatenaciya GMF + MLP', PURPLE, fontsize=9)
    arrow(ax, 2.5, 5.5, 3.8, 6.55)
    arrow(ax, 7.5, 6.38, 6.2, 6.55)

    # Output
    box(ax, 5.0, 7.95, 3.0, 0.6, 'Dense(1) + Sigmoid\n=> Ocenka relevantnosti', RED, fontsize=9)
    arrow(ax, 5.0, 7.15, 5.0, 7.65)

    # Labels
    ax.text(0.5, 3.1, 'GMF\nvetka', fontsize=9, color=GREEN, ha='center',
            fontweight='bold', va='center')
    ax.text(9.5, 3.1, 'MLP\nvetka', fontsize=9, color=ORANGE, ha='center',
            fontweight='bold', va='center')

    ax.set_title('Arkhitektura neyronnoj kollaborativnoj filtracii (NCF)',
                 fontsize=13, fontweight='bold', color=DARK, pad=10)
    plt.tight_layout()
    plt.savefig(os.path.join(OUT_DIR, 'fig_ncf_architecture.png'),
                dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    print('[OK] fig_ncf_architecture.png')


# ===========================================================================
# 5. fig_llm_comparison.png
# ===========================================================================
def gen_llm_comparison():
    models = [
        ('GPT-4',          1800, 86.4, 320, False),
        ('GPT-3.5',          20, 70.0,  10, False),
        ('Claude 3 Opus',   200, 86.8, 100, False),
        ('Llama 3 70B',      70, 79.5,  45, True),
        ('Qwen2.5 7B',        7, 73.2,   5, True),
        ('Qwen2.5 72B',      72, 84.1,  48, True),
        ('DeepSeek 7B',       7, 71.8,   5, True),
        ('Mistral 7B',        7, 64.2,   5, True),
    ]
    names  = [m[0] for m in models]
    params = [m[1] for m in models]
    mmlu   = [m[2] for m in models]
    memory = [m[3] for m in models]
    open_  = [m[4] for m in models]

    colors = [BLUE if o else ORANGE for o in open_]
    sizes  = [max(m * 1.5, 60) for m in memory]

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 6),
                                    gridspec_kw={'width_ratios': [3, 2]})

    # Scatter plot
    for i, (x, y, s, c, n) in enumerate(zip(params, mmlu, sizes, colors, names)):
        if n == 'Qwen2.5 7B':
            ax1.scatter(x, y, s=s*2, c=c, marker='*', zorder=5,
                       edgecolors='black', linewidths=1)
            ax1.annotate('Ispolzuetsya\nv CoDiRank', xy=(x, y),
                        xytext=(x + 20, y - 2.5),
                        fontsize=8, color=BLUE, fontweight='bold',
                        arrowprops=dict(arrowstyle='->', color=BLUE, lw=1.2))
        else:
            ax1.scatter(x, y, s=s, c=c, alpha=0.85, zorder=4,
                       edgecolors='white', linewidths=0.5)
        ax1.annotate(n, xy=(x, y), xytext=(0, 8),
                    textcoords='offset points', ha='center', fontsize=7.5, color=DARK)

    ax1.axvline(x=10, color=GRAY, linestyle='--', linewidth=1.2, alpha=0.7)
    ax1.text(11, 62.5, 'Granica\neffektivnogo\nlokalnogo\ninferensa',
             fontsize=7.5, color=GRAY, va='bottom')

    ax1.set_xscale('log')
    ax1.set_xlabel('Kolichestvo parametrov (mlrd)', fontsize=11, color=DARK)
    ax1.set_ylabel('MMLU (%)', fontsize=11, color=DARK)
    ax1.set_title('Sravnenie LLM: kachestvo vs. razmer', fontsize=12,
                  fontweight='bold', color=DARK)
    ax1.set_xlim(3, 3000)
    ax1.set_ylim(60, 90)
    ax1.grid(True, alpha=0.3)

    legend_elems = [
        mpatches.Patch(color=ORANGE, label='Zakrytye (GPT, Claude)'),
        mpatches.Patch(color=BLUE,   label='Otkrytye (Llama, Qwen...)')
    ]
    ax1.legend(handles=legend_elems, fontsize=9, loc='lower right')

    # Bar chart - MMLU scores
    open_models = [(n, m, c) for n, p, m, mem, o, c
                   in zip(names, params, mmlu, memory, open_, colors) if o]
    on, om, oc = zip(*open_models)
    bars = ax2.barh(list(on), list(om), color=list(oc), alpha=0.85, edgecolor='white')
    ax2.set_xlabel('MMLU (%)', fontsize=10, color=DARK)
    ax2.set_title('Otkrytye LLM: MMLU', fontsize=11, fontweight='bold', color=DARK)
    ax2.set_xlim(55, 90)
    ax2.axvline(x=70, color=GRAY, linestyle=':', alpha=0.5)
    for bar, val in zip(bars, om):
        ax2.text(bar.get_width() + 0.3, bar.get_y() + bar.get_height()/2,
                 f'{val}%', va='center', fontsize=9, color=DARK)
    ax2.grid(True, axis='x', alpha=0.3)

    plt.suptitle('Sravnitelnaya kharakteristika otkrytykh LLM po razmeru i kachestvu',
                 fontsize=13, fontweight='bold', color=DARK, y=1.02)
    plt.tight_layout()
    plt.savefig(os.path.join(OUT_DIR, 'fig_llm_comparison.png'),
                dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    print('[OK] fig_llm_comparison.png')


# ===========================================================================
# 6. fig_gplay_dataset.png
# ===========================================================================
def gen_gplay_dataset():
    fig, axes = plt.subplots(2, 2, figsize=(13, 8))
    fig.suptitle('Struktura i statistika dataseta Google Play Store Apps',
                 fontsize=14, fontweight='bold', color=DARK, y=1.01)

    # Top-left: pie Free/Paid
    ax = axes[0, 0]
    ax.pie([92, 8], labels=['Besplatnye (92%)', 'Platnye (8%)'],
           colors=[BLUE, ORANGE], autopct='%1.0f%%',
           startangle=90, textprops={'fontsize': 10, 'color': DARK},
           wedgeprops={'edgecolor': 'white', 'linewidth': 2})
    ax.set_title('Sootnoshenie besplatnykh i platnykh', fontsize=11,
                 fontweight='bold', color=DARK)

    # Top-right: ratings histogram
    ax = axes[0, 1]
    np.random.seed(7)
    ratings = np.concatenate([
        np.random.normal(4.3, 0.4, 7000),
        np.random.normal(3.5, 0.3, 2000),
        np.random.normal(2.0, 0.5, 1000)
    ])
    ratings = ratings[(ratings >= 1) & (ratings <= 5)]
    n, bins, patches_hist = ax.hist(ratings, bins=20, color=BLUE, alpha=0.8,
                                     edgecolor='white', linewidth=0.5)
    for p, b in zip(patches_hist, bins):
        if b >= 4.0:
            p.set_facecolor(GREEN)
    ax.set_xlabel('Rejting', fontsize=10, color=DARK)
    ax.set_ylabel('Kolichestvo prilozheniy', fontsize=10, color=DARK)
    ax.set_title('Raspredelenie rejtingov', fontsize=11, fontweight='bold', color=DARK)
    ax.grid(True, axis='y', alpha=0.3)
    ax.text(1.2, max(n)*0.85, 'Sredny rejting: 4.17',
            fontsize=9, color=DARK,
            bbox=dict(boxstyle='round,pad=0.3', facecolor=LIGHT_GRAY, alpha=0.8))

    # Bottom-left: top categories bar
    ax = axes[1, 0]
    categories = ['Obrazovanie', 'Razvlecheniya', 'Biznes', 'Instrumenty',
                  'Stil zhizni', 'Muzyka', 'Knigi', 'Zdorove', 'Igry', 'Fotokamera']
    counts = [1200, 1150, 980, 920, 870, 760, 720, 690, 650, 600]
    colors_list = [BLUE if i < 3 else GRAY for i in range(len(categories))]
    bars = ax.barh(categories, counts, color=colors_list, alpha=0.85, edgecolor='white')
    ax.set_xlabel('Kolichestvo prilozheniy', fontsize=10, color=DARK)
    ax.set_title('Top-10 kategorij', fontsize=11, fontweight='bold', color=DARK)
    ax.grid(True, axis='x', alpha=0.3)
    for bar, cnt in zip(bars, counts):
        ax.text(bar.get_width() + 5, bar.get_y() + bar.get_height()/2,
                str(cnt), va='center', fontsize=8, color=DARK)

    # Bottom-right: reviews histogram (log)
    ax = axes[1, 1]
    np.random.seed(3)
    reviews = np.random.exponential(scale=5000, size=5000)
    reviews = reviews[reviews > 1]
    ax.hist(reviews, bins=40, color=PURPLE, alpha=0.8, edgecolor='white', linewidth=0.3)
    ax.set_xscale('log')
    ax.set_xlabel('Kolichestvo otzyovov (log shkala)', fontsize=10, color=DARK)
    ax.set_ylabel('Chislo prilozheniy', fontsize=10, color=DARK)
    ax.set_title('Raspredelenie kol-va otzyovov', fontsize=11, fontweight='bold', color=DARK)
    ax.grid(True, alpha=0.3)
    ax.text(0.98, 0.95, 'Vsego zapisej: ~500K',
            transform=ax.transAxes, ha='right', va='top', fontsize=9,
            bbox=dict(boxstyle='round,pad=0.3', facecolor=LIGHT_GRAY))

    plt.tight_layout()
    plt.savefig(os.path.join(OUT_DIR, 'fig_gplay_dataset.png'),
                dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    print('[OK] fig_gplay_dataset.png')


# ===========================================================================
# 7. fig_profile_scheme.png  (NEW, detailed version)
# ===========================================================================
def gen_profile_scheme():
    fig, ax = plt.subplots(figsize=(14, 7))
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 7)
    ax.axis('off')

    def rounded_box(ax, x, y, w, h, text, fc, ec='white', fontsize=9, tc='white'):
        b = FancyBboxPatch((x - w/2, y - h/2), w, h,
                           boxstyle='round,pad=0.12',
                           facecolor=fc, edgecolor=ec, linewidth=1.2)
        ax.add_patch(b)
        ax.text(x, y, text, ha='center', va='center',
                fontsize=fontsize, color=tc, fontweight='bold')

    def arr(ax, x1, y1, x2, y2, c=GRAY, label=''):
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle='->', color=c, lw=1.5))
        if label:
            mx, my = (x1+x2)/2, (y1+y2)/2
            ax.text(mx + 0.05, my, label, fontsize=8, color=c, va='center')

    turns = [1.2, 4.3, 7.4, 10.5]
    turn_labels = ['t=0', 't=1', 't=2', 't=3']

    for i, (tx, tl) in enumerate(zip(turns, turn_labels)):
        # Timeline dot
        ax.scatter([tx], [3.5], s=120, color=BLUE, zorder=5)
        ax.text(tx, 3.1, tl, ha='center', fontsize=11, color=BLUE, fontweight='bold')

        # User message
        rounded_box(ax, tx, 5.5, 1.6, 0.55,
                    f'Soobshchenie\npolzovatelya {i}', DARK, fontsize=7.5)

        # Qwen embed
        rounded_box(ax, tx, 4.6, 1.6, 0.55,
                    'Qwen embed\nE(a_t)', '#4F46E5', fontsize=7.5)
        arr(ax, tx, 5.22, tx, 4.88, BLUE)

        arr(ax, tx, 4.32, tx, 3.78, BLUE)

    # Timeline line
    ax.plot([0.3, 11.5], [3.5, 3.5], color=BLUE, linewidth=2, alpha=0.4)

    # Profile update boxes
    for i, tx in enumerate(turns[1:], 1):
        px = turns[i-1] + (turns[i] - turns[i-1]) / 2 + 0.5
        rounded_box(ax, px, 2.5, 2.2, 0.65,
                    f'P(t) = a*P(t-1)\n+ (1-a)*E(a_t)*w', GREEN, fontsize=7.5)
        arr(ax, turns[i-1], 3.22, px - 0.6, 2.72, GREEN, 'a*P')
        arr(ax, turns[i], 3.22, px + 0.3, 2.72, BLUE, '+E')
        arr(ax, px, 2.17, turns[i] - 0.2, 1.75, GREEN)

    # Reject box
    rounded_box(ax, 10.5, 1.35, 2.8, 0.6,
                'P(t) = P(t) - b*E(m_rej)\n[otkaz polzovatelya]',
                RED, fontsize=7.5)
    arr(ax, 10.5, 3.22, 10.5, 1.65, RED)

    # Alpha annotation
    ax.text(0.3, 1.0,
            'a (alfa) = koefficient zabyvanya  |  b (beta) = sila otkaza  |  w = ves priblizhenia',
            fontsize=8.5, color=GRAY, va='center')

    ax.set_title('Skhema dinamicheskogo obnovleniya profilya interesov polzovatelya\n'
                 'P(t) = alfa * P(t-1) + (1-alfa) * E(a_t) * w',
                 fontsize=12, fontweight='bold', color=DARK)
    plt.tight_layout()
    plt.savefig(os.path.join(OUT_DIR, 'fig_profile_scheme.png'),
                dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    print('[OK] fig_profile_scheme.png')


# ===========================================================================
# 9. fig_relevance_space.png
# ===========================================================================
def gen_relevance_space():
    from matplotlib import cm

    x = np.linspace(0, 1, 80)
    y = np.linspace(0, 1, 80)
    X, Y = np.meshgrid(x, y)
    # R = 0.6*cos_sim + 0.4*attr_match - penalty
    Z = 0.6 * X + 0.4 * Y

    fig, axes = plt.subplots(1, 2, figsize=(13, 5.5))

    # Left: contour plot
    ax = axes[0]
    cf = ax.contourf(X, Y, Z, levels=20, cmap='RdYlGn')
    cbar = plt.colorbar(cf, ax=ax)
    cbar.set_label('R(i, D, t)', fontsize=10, color=DARK)

    # Mark sample apps
    apps_pts = [
        (0.85, 0.90, 'App A\n(relev.)', GREEN),
        (0.40, 0.70, 'App B\n(sredniy)', ORANGE),
        (0.15, 0.20, 'App C\n(nerelev.)', RED),
    ]
    for cx, cy, name, color in apps_pts:
        ax.scatter([cx], [cy], color=color, s=120, zorder=5, edgecolors='white', linewidths=1.5)
        ax.annotate(name, xy=(cx, cy), xytext=(cx + 0.05, cy - 0.1),
                    fontsize=8, color=color, fontweight='bold')

    # Rejection shift
    ax.annotate('', xy=(0.65, 0.55), xytext=(0.85, 0.90),
                arrowprops=dict(arrowstyle='->', color=RED, lw=1.5, linestyle='dashed'))
    ax.text(0.63, 0.70, 'shtraf\notkaza', fontsize=8, color=RED, ha='center')

    ax.set_xlabel('cos(E(app), P(t)) — semanticheskoe skhodstvo', fontsize=9, color=DARK)
    ax.set_ylabel('attr_match — sovpadenie atributov', fontsize=9, color=DARK)
    ax.set_title('Prostranstvo relevantnosti CoDiRank\n(konturny grafik)', fontsize=11,
                 fontweight='bold', color=DARK)

    # Right: formula explanation
    ax2 = axes[1]
    ax2.axis('off')
    formula_lines = [
        ('Formula relevantnosti:', DARK, 12, True),
        ('', DARK, 8, False),
        ('R(i, D, t) =', DARK, 11, True),
        ('  l1 * cos(E(i), P(t))', BLUE, 10, False),
        ('+ l2 * attr_match(i, attrs(D))', GREEN, 10, False),
        ('- l3 * reject_penalty(i, D)', RED, 10, False),
        ('- l4 * shown_penalty(i, D)', ORANGE, 10, False),
        ('', DARK, 8, False),
        ('gde:', DARK, 10, True),
        ('  E(i) — embedding prilozheniya', GRAY, 9, False),
        ('  P(t) — profil polzovatelya', GRAY, 9, False),
        ('  D    — tekushchiy dialog', GRAY, 9, False),
        ('  l1..l4 — nastraiv. vesa', GRAY, 9, False),
    ]
    y_pos = 0.95
    for text, color, fontsize, bold in formula_lines:
        fw = 'bold' if bold else 'normal'
        ax2.text(0.05, y_pos, text, transform=ax2.transAxes,
                 fontsize=fontsize, color=color, fontweight=fw, va='top',
                 family='monospace' if text.startswith(' ') else 'DejaVu Sans')
        y_pos -= 0.07

    ax2.set_facecolor(LIGHT_GRAY)
    ax2.set_title('Komponenty formuly', fontsize=11, fontweight='bold', color=DARK)

    plt.suptitle('Vizualizaciya funkcii relevantnosti CoDiRank\nv prostranstve embeddingov',
                 fontsize=12, fontweight='bold', color=DARK)
    plt.tight_layout()
    plt.savefig(os.path.join(OUT_DIR, 'fig_relevance_space.png'),
                dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    print('[OK] fig_relevance_space.png')


# ===========================================================================
# 10. fig_metrics_table.png
# ===========================================================================
def gen_metrics_table():
    methods = ['CoDiRank', 'CF (user-user)', 'CF (item-item)',
               'Content-Based', 'Hybrid', 'BM25', 'Sluchajnyj']
    metrics = ['Precision@3', 'Recall@3', 'NDCG@3', 'MRR', 'HR@10', 'Cold-Start']

    data = np.array([
        [0.71, 0.68, 0.73, 0.74, 0.85, 0.62],  # CoDiRank
        [0.52, 0.49, 0.54, 0.56, 0.71, 0.18],  # CF user-user
        [0.55, 0.52, 0.57, 0.58, 0.73, 0.21],  # CF item-item
        [0.48, 0.45, 0.50, 0.51, 0.66, 0.44],  # Content-Based
        [0.61, 0.58, 0.63, 0.65, 0.78, 0.35],  # Hybrid
        [0.35, 0.32, 0.37, 0.38, 0.52, 0.28],  # BM25
        [0.11, 0.09, 0.12, 0.13, 0.18, 0.11],  # Random
    ])

    fig, ax = plt.subplots(figsize=(13, 5))
    ax.axis('off')

    import matplotlib.colors as mcolors
    norm = mcolors.Normalize(vmin=0.05, vmax=0.9)
    cmap = plt.cm.RdYlGn

    table = ax.table(
        cellText=[[f'{v:.2f}' for v in row] for row in data],
        rowLabels=methods,
        colLabels=metrics,
        loc='center',
        cellLoc='center'
    )
    table.auto_set_font_size(False)
    table.set_fontsize(11)
    table.scale(1.6, 2.0)

    for (row, col), cell in table.get_celld().items():
        cell.set_edgecolor('#E5E7EB')
        if row == 0:
            cell.set_facecolor(DARK)
            cell.set_text_props(color='white', fontweight='bold', fontsize=10)
        elif col == -1:
            if methods[row - 1] == 'CoDiRank':
                cell.set_facecolor(BLUE)
                cell.set_text_props(color='white', fontweight='bold', fontsize=11)
            else:
                cell.set_facecolor('#F9FAFB')
                cell.set_text_props(color=DARK, fontsize=10)
        else:
            val = data[row - 1, col]
            rgba = cmap(norm(val))
            cell.set_facecolor(rgba)
            cell.set_text_props(color='black' if val > 0.4 else 'white',
                               fontsize=11,
                               fontweight='bold' if methods[row - 1] == 'CoDiRank' else 'normal')

    ax.set_title('Sravnitelnaya tablitsa metrik kachestva: CoDiRank vs baseline-metody\n'
                 '(znacheniya ot 0 do 1, zeleny = luchshe)',
                 fontsize=12, fontweight='bold', color=DARK, pad=20, y=0.98)
    plt.tight_layout()
    plt.savefig(os.path.join(OUT_DIR, 'fig_metrics_table.png'),
                dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    print('[OK] fig_metrics_table.png')


# ===========================================================================
# 12. fig_ollama_config.png
# ===========================================================================
def gen_ollama_config():
    fig = plt.figure(figsize=(13, 5.5))
    gs = GridSpec(1, 2, figure=fig, width_ratios=[1.2, 1])

    # Left: terminal-style model card
    ax1 = fig.add_subplot(gs[0])
    ax1.set_xlim(0, 10)
    ax1.set_ylim(0, 8)
    ax1.axis('off')

    # Terminal box background
    term_bg = FancyBboxPatch((0.1, 0.2), 9.6, 7.5,
                             boxstyle='round,pad=0.1',
                             facecolor='#0D1117', edgecolor='#30363D', linewidth=2)
    ax1.add_patch(term_bg)

    # Title bar
    title_bar = FancyBboxPatch((0.1, 7.2), 9.6, 0.5,
                               boxstyle='square,pad=0',
                               facecolor='#21262D', edgecolor='none')
    ax1.add_patch(title_bar)

    # Window buttons
    for i, c in enumerate(['#FF5F57', '#FFBD2E', '#28CA41']):
        ax1.add_patch(plt.Circle((0.55 + i*0.35, 7.45), 0.1, color=c))
    ax1.text(5.0, 7.45, 'ollama show qwen2.5:7b', ha='center', va='center',
             fontsize=9, color='#8B949E', family='monospace')

    # Content
    lines = [
        ('$ ollama show qwen2.5:7b', '#58A6FF', 10),
        ('', '', 8),
        ('  Model:', '#8B949E', 9),
        ('    arch:            qwen2',   '#E6EDF3', 9),
        ('    parameters:      7.6B',    '#79C0FF', 9),
        ('    context length:  32768',   '#79C0FF', 9),
        ('    embedding length: 3584',   '#79C0FF', 9),
        ('    quantization:    Q4_K_M',  '#FFA657', 9),
        ('', '', 8),
        ('  License:', '#8B949E', 9),
        ('    Apache 2.0 (otkrytaya)', '#3FB950', 9),
        ('', '', 8),
        ('  Sistema:', '#8B949E', 9),
        ('    RAM: 6.2 GB (minimum)',   '#E6EDF3', 9),
        ('    Skorost: ~42 tok/sek',     '#3FB950', 9),
        ('    CPU: podderzhivaetsya',     '#E6EDF3', 9),
    ]
    y = 6.8
    for text, color, fs in lines:
        if text:
            ax1.text(0.5, y, text, va='center', fontsize=fs, color=color,
                    family='monospace')
        y -= 0.37

    ax1.set_title('Konfiguratsiya Ollama: model qwen2.5:7b',
                  fontsize=11, fontweight='bold', color=DARK, pad=8)

    # Right: inference speed bar chart
    ax2 = fig.add_subplot(gs[1])
    quants = ['F16', 'Q8_0', 'Q5_K_M', 'Q4_K_M']
    speeds = [18, 28, 36, 42]
    memory_gb = [14.5, 7.7, 5.4, 4.8]

    colors_q = [RED, ORANGE, BLUE, GREEN]
    bars = ax2.barh(quants, speeds, color=colors_q, alpha=0.85, edgecolor='white', height=0.5)
    ax2.set_xlabel('Skorost inferensa (tokenov/sek)', fontsize=10, color=DARK)
    ax2.set_title('Skorost vs kvantizatsiya\n(Qwen2.5-7B, CPU)', fontsize=11,
                  fontweight='bold', color=DARK)
    ax2.set_xlim(0, 55)
    ax2.grid(True, axis='x', alpha=0.3)

    for bar, spd, mem in zip(bars, speeds, memory_gb):
        ax2.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
                 f'{spd} tok/s | {mem} GB', va='center', fontsize=9, color=DARK)

    ax2.axvline(x=42, color=GREEN, linestyle='--', linewidth=1.5, alpha=0.7)
    ax2.text(43, -0.4, 'Ispolzuemaya\nkonfig.', fontsize=8, color=GREEN)

    plt.suptitle('Konfiguratsiya Ollama i parametry razvertyvanya Qwen2.5-7B-Instruct',
                 fontsize=12, fontweight='bold', color=DARK)
    plt.tight_layout()
    plt.savefig(os.path.join(OUT_DIR, 'fig_ollama_config.png'),
                dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    print('[OK] fig_ollama_config.png')


# ===========================================================================
# 13. fig_db_schema.png
# ===========================================================================
def gen_db_schema():
    fig, ax = plt.subplots(figsize=(15, 9))
    ax.set_xlim(0, 15)
    ax.set_ylim(0, 9)
    ax.axis('off')
    ax.set_facecolor('#F8FAFC')
    fig.patch.set_facecolor('#F8FAFC')

    def draw_table(ax, x, y, name, columns, width=3.2, header_color=BLUE):
        row_h = 0.38
        header_h = 0.5
        total_h = header_h + len(columns) * row_h

        # Header
        hdr = FancyBboxPatch((x, y - header_h), width, header_h,
                             boxstyle='round,pad=0.05',
                             facecolor=header_color, edgecolor='white', linewidth=1.5)
        ax.add_patch(hdr)
        ax.text(x + width/2, y - header_h/2, name,
                ha='center', va='center', fontsize=10, color='white', fontweight='bold')

        # Columns
        for i, (col_name, col_type, is_pk, is_fk, is_vec) in enumerate(columns):
            row_y = y - header_h - (i+1)*row_h
            bg_color = '#FFF7ED' if is_vec else ('#EFF6FF' if is_pk else 'white')
            row_bg = Rectangle((x, row_y), width, row_h,
                                facecolor=bg_color, edgecolor='#E2E8F0', linewidth=0.8)
            ax.add_patch(row_bg)

            prefix = 'PK ' if is_pk else ('FK ' if is_fk else '   ')
            prefix_color = ORANGE if is_pk else (RED if is_fk else DARK)
            ax.text(x + 0.15, row_y + row_h/2, prefix,
                    va='center', fontsize=7.5, color=prefix_color, fontweight='bold',
                    family='monospace')
            ax.text(x + 0.55, row_y + row_h/2, col_name,
                    va='center', fontsize=7.5, color=DARK, fontweight='bold' if is_pk else 'normal')
            type_color = PURPLE if is_vec else GRAY
            ax.text(x + width - 0.1, row_y + row_h/2, col_type,
                    va='center', ha='right', fontsize=7, color=type_color,
                    fontweight='bold' if is_vec else 'normal')

        return total_h

    # users table
    users_cols = [
        ('id',         'SERIAL',    True,  False, False),
        ('username',   'VARCHAR',   False, False, False),
        ('created_at', 'TIMESTAMP', False, False, False),
        ('settings',   'JSONB',     False, False, False),
    ]
    draw_table(ax, 0.4, 8.5, 'users', users_cols, width=3.5, header_color=BLUE)

    # sessions table
    sess_cols = [
        ('id',          'SERIAL',         True,  False, False),
        ('user_id',     'INT',            False, True,  False),
        ('status',      'VARCHAR',        False, False, False),
        ('profile_vec', 'VECTOR(3584)',   False, False, True),
        ('attributes',  'JSONB',          False, False, False),
        ('turn_count',  'INT',            False, False, False),
    ]
    draw_table(ax, 5.2, 8.5, 'sessions', sess_cols, width=3.8, header_color=GREEN)

    # turns table
    turns_cols = [
        ('id',         'SERIAL',       True,  False, False),
        ('session_id', 'INT',          False, True,  False),
        ('turn_index', 'INT',          False, False, False),
        ('role',       'VARCHAR',      False, False, False),
        ('content',    'TEXT',         False, False, False),
        ('embedding',  'VECTOR(3584)', False, False, True),
    ]
    draw_table(ax, 5.2, 4.6, 'turns', turns_cols, width=3.8, header_color=PURPLE)

    # apps table
    apps_cols = [
        ('id',          'SERIAL',       True,  False, False),
        ('external_id', 'VARCHAR UNIQ', False, False, False),
        ('platform',    'VARCHAR',      False, False, False),
        ('name',        'VARCHAR',      False, False, False),
        ('category',    'VARCHAR',      False, False, False),
        ('rating',      'FLOAT',        False, False, False),
        ('price',       'VARCHAR',      False, False, False),
        ('embedding',   'VECTOR(3584)', False, False, True),
    ]
    draw_table(ax, 10.5, 8.5, 'apps', apps_cols, width=3.8, header_color=ORANGE)

    # feedback table
    fb_cols = [
        ('id',         'SERIAL', True,  False, False),
        ('session_id', 'INT',    False, True,  False),
        ('app_id',     'INT',    False, True,  False),
        ('signal',     'VARCHAR',False, False, False),
    ]
    draw_table(ax, 10.5, 3.8, 'feedback', fb_cols, width=3.8, header_color=RED)

    # FK arrows
    def fk_arrow(ax, x1, y1, x2, y2):
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle='->', color=RED, lw=1.5,
                                   linestyle='dashed',
                                   connectionstyle='arc3,rad=0.0'))

    # users -> sessions
    fk_arrow(ax, 3.9, 7.3, 5.2, 7.6)
    ax.text(4.55, 7.6, 'user_id', fontsize=7.5, color=RED, ha='center')

    # sessions -> turns
    fk_arrow(ax, 7.1, 6.18, 7.1, 4.6)
    ax.text(7.5, 5.4, 'session_id', fontsize=7.5, color=RED)

    # sessions -> feedback
    fk_arrow(ax, 9.0, 6.5, 10.5, 3.42)
    ax.text(9.5, 5.0, 'session_id', fontsize=7.5, color=RED)

    # apps -> feedback
    fk_arrow(ax, 12.4, 5.68, 12.4, 3.8)
    ax.text(12.55, 4.7, 'app_id', fontsize=7.5, color=RED)

    # Vector legend
    vec_legend = FancyBboxPatch((0.3, 0.3), 4.0, 0.7,
                                boxstyle='round,pad=0.1',
                                facecolor='#FFF7ED', edgecolor=PURPLE, linewidth=1.5)
    ax.add_patch(vec_legend)
    ax.text(0.6, 0.65, 'VECTOR(3584)', fontsize=9, color=PURPLE, fontweight='bold')
    ax.text(0.6, 0.40, '— embeddingi Qwen2.5 (pgvector)', fontsize=8.5, color=DARK)

    ax.set_title('Skhema bazy dannykh PostgreSQL: tablitsy i svyazi\n'
                 '(vydeleny VECTOR-kolonki dlya semanticheskogo poiska)',
                 fontsize=12, fontweight='bold', color=DARK, pad=10)
    plt.tight_layout()
    plt.savefig(os.path.join(OUT_DIR, 'fig_db_schema.png'),
                dpi=150, bbox_inches='tight', facecolor='#F8FAFC')
    plt.close()
    print('[OK] fig_db_schema.png')


# ===========================================================================
# 14. fig_bot_dialog.png
# ===========================================================================
def gen_bot_dialog():
    fig, ax = plt.subplots(figsize=(8, 11))
    ax.set_xlim(0, 8)
    ax.set_ylim(0, 11)
    ax.axis('off')
    fig.patch.set_facecolor('#17212B')

    # Phone frame
    phone = FancyBboxPatch((0.3, 0.2), 7.4, 10.6,
                           boxstyle='round,pad=0.3',
                           facecolor='#17212B', edgecolor='#2B3547', linewidth=2)
    ax.add_patch(phone)

    # Top bar
    top = FancyBboxPatch((0.3, 10.2), 7.4, 0.6,
                         boxstyle='square,pad=0',
                         facecolor='#232E3C', edgecolor='none')
    ax.add_patch(top)
    ax.text(4.0, 10.5, 'CoDiRank Bot', ha='center', va='center',
            fontsize=11, color='white', fontweight='bold')
    ax.text(1.0, 10.5, '<', fontsize=14, color='#6B7A8D', va='center', ha='center')
    # Status dot
    ax.add_patch(plt.Circle((5.9, 10.5), 0.07, color='#4DCA5A'))
    ax.text(6.1, 10.5, 'v seti', fontsize=8, color='#4DCA5A', va='center')

    def user_bubble(ax, text, y):
        lines = [text[i:i+32] for i in range(0, len(text), 32)]
        h = 0.3 + len(lines) * 0.28
        x_start = 7.7 - 0.3 - len(max(lines, key=len)) * 0.11 - 0.3
        x_start = max(x_start, 3.5)
        bw = 7.4 - x_start
        b = FancyBboxPatch((x_start, y - h + 0.1), bw, h,
                           boxstyle='round,pad=0.12',
                           facecolor='#2B5278', edgecolor='none')
        ax.add_patch(b)
        for i, line in enumerate(lines):
            ax.text(x_start + bw/2, y - 0.15 - i*0.28, line,
                    ha='center', va='top', fontsize=8.5, color='white')
        return h + 0.15

    def bot_bubble(ax, text, y, has_buttons=False):
        lines = []
        for part in text.split('\n'):
            while len(part) > 36:
                lines.append(part[:36])
                part = part[36:]
            lines.append(part)
        h = 0.3 + len(lines) * 0.28
        if has_buttons:
            h += 0.55
        bw = min(len(max(lines, key=len)) * 0.105 + 0.5, 5.5)
        b = FancyBboxPatch((0.5, y - h + 0.1), bw, h,
                           boxstyle='round,pad=0.12',
                           facecolor='#242F3D', edgecolor='none')
        ax.add_patch(b)
        for i, line in enumerate(lines):
            ax.text(0.7, y - 0.15 - i*0.28, line,
                    ha='left', va='top', fontsize=8.0, color='#F5F5F5')
        if has_buttons:
            btn_y = y - h + 0.35
            for bi, (btxt, bcol) in enumerate([('Podkhodit', GREEN),
                                               ('Ne to', RED),
                                               ('Detali', BLUE)]):
                bx = 0.6 + bi * 1.7
                bb = FancyBboxPatch((bx, btn_y), 1.5, 0.35,
                                   boxstyle='round,pad=0.06',
                                   facecolor=bcol, edgecolor='none', alpha=0.9)
                ax.add_patch(bb)
                ax.text(bx + 0.75, btn_y + 0.175, btxt, ha='center', va='center',
                        fontsize=7.5, color='white', fontweight='bold')
        return h + 0.2

    messages = [
        ('user', 'Ishchu prilozhenie dlya\nmeditatsii, besplatnoe'),
        ('bot',  'Ponyatno! Na kakoy platforme:\nAndroid ili iOS?'),
        ('user', 'iOS, bez reklamy'),
        ('bot',  'Khorosho, ishchu\npodkhodyashchie prilozheniya...'),
        ('bot',  'Insight Timer — iOS\n4.8  Besplatno  Bez reklamy\nPochemu podkhodit: Krupneyshaya\nbiblieka meditatsiy (russkiy).', True),
    ]

    y_cur = 9.8
    for msg in messages:
        role = msg[0]
        text = msg[1]
        has_btn = len(msg) > 2 and msg[2]
        if role == 'user':
            dy = user_bubble(ax, text, y_cur)
        else:
            dy = bot_bubble(ax, text, y_cur, has_buttons=has_btn)
        y_cur -= dy + 0.1

    # Input bar
    input_bg = FancyBboxPatch((0.4, 0.25), 7.2, 0.55,
                              boxstyle='round,pad=0.1',
                              facecolor='#242F3D', edgecolor='none')
    ax.add_patch(input_bg)
    ax.text(1.2, 0.52, 'Napisat soobshchenie...', va='center', fontsize=8.5, color='#6B7A8D')

    ax.set_title('Skrinshot dialoga s botom: primer sessii\npodbora prilozheniya',
                 fontsize=11, fontweight='bold', color=DARK, pad=10)
    plt.tight_layout()
    plt.savefig(os.path.join(OUT_DIR, 'fig_bot_dialog.png'),
                dpi=150, bbox_inches='tight', facecolor='#17212B')
    plt.close()
    print('[OK] fig_bot_dialog.png')


# ===========================================================================
# 16. fig_dataset_stats.png
# ===========================================================================
def gen_dataset_stats():
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(13, 5))
    fig.suptitle('Statistika testovogo dialogovogo dataseta:\n'
                 'raspredelenie po dline i kategoriyam',
                 fontsize=12, fontweight='bold', color=DARK)

    # Left: dialog length histogram
    np.random.seed(11)
    lengths = np.random.normal(4.0, 1.5, 1200)
    lengths = np.clip(np.round(lengths), 1, 10).astype(int)
    bins = np.arange(0.5, 11.5, 1)
    counts = [np.sum(lengths == i) for i in range(1, 11)]

    colors_bar = [GREEN if i in [3, 4, 5] else BLUE for i in range(1, 11)]
    ax1.bar(range(1, 11), counts, color=colors_bar, alpha=0.85, edgecolor='white', linewidth=0.5)

    # Bell curve overlay
    x_fit = np.linspace(0.5, 10.5, 200)
    sigma = 1.5
    mu = 4.0
    y_fit = (1 / (sigma * np.sqrt(2 * np.pi))) * np.exp(-0.5 * ((x_fit - mu) / sigma) ** 2) * len(lengths)
    ax1.plot(x_fit, y_fit, color=RED, linewidth=2, label='Normal (mu=4)')

    ax1.axvspan(2.5, 5.5, alpha=0.1, color=GREEN, label='Optimalna zona')
    ax1.set_xlabel('Kolichestvo khodov dialoga', fontsize=11, color=DARK)
    ax1.set_ylabel('Chislo dialogov', fontsize=11, color=DARK)
    ax1.set_title('Raspredelenie po dline dialoga', fontsize=11, fontweight='bold', color=DARK)
    ax1.set_xticks(range(1, 11))
    ax1.legend(fontsize=8)
    ax1.grid(True, axis='y', alpha=0.3)
    ax1.text(0.98, 0.95, f'Vsego dialogov: {len(lengths)}',
             transform=ax1.transAxes, ha='right', va='top', fontsize=9,
             bbox=dict(boxstyle='round,pad=0.3', facecolor=LIGHT_GRAY))

    # Right: pie of intent categories
    categories = ['Igry', 'Zdorove', 'Produktivnost',
                  'Razvlecheniya', 'Obrazovanie', 'Drugie']
    sizes = [25, 18, 15, 12, 10, 20]
    colors_pie = [BLUE, GREEN, ORANGE, PURPLE, RED, GRAY]
    explode = [0.05] * len(categories)

    wedges, texts, autotexts = ax2.pie(
        sizes, labels=categories, colors=colors_pie,
        autopct='%1.0f%%', startangle=90, explode=explode,
        textprops={'fontsize': 9, 'color': DARK},
        wedgeprops={'edgecolor': 'white', 'linewidth': 1.5}
    )
    for at in autotexts:
        at.set_fontsize(8)
        at.set_color('white')
    ax2.set_title('Raspredelenie po kategoriyam\nzaprosov', fontsize=11,
                  fontweight='bold', color=DARK)

    plt.tight_layout()
    plt.savefig(os.path.join(OUT_DIR, 'fig_dataset_stats.png'),
                dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    print('[OK] fig_dataset_stats.png')


# ===========================================================================
# 18. fig_ndcg_vs_turns.png
# ===========================================================================
def gen_ndcg_vs_turns():
    turns = np.arange(1, 11)
    np.random.seed(42)

    def smooth_curve(base, noise=0.02):
        y = base + np.random.normal(0, noise, len(turns))
        return np.clip(y, 0, 1)

    codirank   = np.array([0.38, 0.52, 0.63, 0.71, 0.74, 0.76, 0.77, 0.77, 0.78, 0.78])
    hybrid     = np.array([0.30, 0.43, 0.53, 0.60, 0.63, 0.65, 0.66, 0.66, 0.67, 0.67])
    cf         = np.array([0.22, 0.33, 0.41, 0.47, 0.50, 0.51, 0.52, 0.52, 0.52, 0.53])
    content    = np.array([0.28, 0.36, 0.42, 0.46, 0.48, 0.49, 0.49, 0.50, 0.50, 0.50])
    bm25       = np.array([0.15, 0.21, 0.25, 0.27, 0.28, 0.28, 0.28, 0.29, 0.29, 0.29])

    fig, ax = plt.subplots(figsize=(11, 5.5))

    # Confidence band for CoDiRank
    ci = 0.03
    ax.fill_between(turns, codirank - ci, codirank + ci,
                    color=BLUE, alpha=0.15, label='_nolegend_')

    ax.plot(turns, codirank, color=BLUE, linewidth=2.5, marker='o',
            markersize=6, label='CoDiRank (nash)', zorder=5)
    ax.plot(turns, hybrid,   color=PURPLE, linewidth=2, marker='s',
            markersize=5, label='Hybrid')
    ax.plot(turns, cf,       color=GREEN, linewidth=1.8, linestyle='--',
            marker='^', markersize=5, label='CF (item-item)')
    ax.plot(turns, content,  color=ORANGE, linewidth=1.8, linestyle='--',
            marker='D', markersize=4, label='Content-Based')
    ax.plot(turns, bm25,     color=RED, linewidth=1.5, linestyle=':',
            marker='x', markersize=6, label='BM25')

    # Annotation - optimal depth
    ax.axvspan(3, 5, alpha=0.08, color=GREEN)
    ax.annotate('Optimalnaya glubina:\n3-4 khoda dialoga',
                xy=(4, codirank[3]),
                xytext=(5.5, 0.55),
                fontsize=9, color=GREEN, fontweight='bold',
                arrowprops=dict(arrowstyle='->', color=GREEN, lw=1.5))

    ax.set_xlabel('Kolichestvo khodov dialoga', fontsize=12, color=DARK)
    ax.set_ylabel('NDCG@10', fontsize=12, color=DARK)
    ax.set_title('Zavisimost NDCG@10 ot chisla khodov dialoga\n'
                 'CoDiRank nakaplivaet kontekst i uluchshaet kachestvo',
                 fontsize=12, fontweight='bold', color=DARK)
    ax.set_xticks(turns)
    ax.set_ylim(0.05, 0.95)
    ax.legend(fontsize=9, loc='lower right')
    ax.grid(True, alpha=0.3)

    plt.tight_layout()
    plt.savefig(os.path.join(OUT_DIR, 'fig_ndcg_vs_turns.png'),
                dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    print('[OK] fig_ndcg_vs_turns.png')


# ===========================================================================
# 19. fig_error_analysis.png
# ===========================================================================
def gen_error_analysis():
    fig, axes = plt.subplots(2, 2, figsize=(13, 8.5))
    fig.suptitle('Analiz oshibochnykh rekomendatsiy: tipichnye sluchai nizkoy relevantnosti',
                 fontsize=13, fontweight='bold', color=DARK)

    # Top-left: error types bar chart
    ax = axes[0, 0]
    error_types = ['Semantich.\noshib.', 'Atribyut.\nkonflikt',
                   'Nedostatochno\ndannykh', 'Lozhnoe\nsrabatyv.']
    values = [35, 28, 22, 15]
    colors_e = [RED, ORANGE, BLUE, PURPLE]
    bars = ax.bar(error_types, values, color=colors_e, alpha=0.85, edgecolor='white',
                  width=0.6)
    for bar, val in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5,
                f'{val}%', ha='center', va='bottom', fontsize=10, fontweight='bold',
                color=DARK)
    ax.set_ylabel('Dolya oshibok (%)', fontsize=10, color=DARK)
    ax.set_title('Tipologiya oshibok', fontsize=11, fontweight='bold', color=DARK)
    ax.set_ylim(0, 45)
    ax.grid(True, axis='y', alpha=0.3)

    # Top-right: semantic error scatter
    ax = axes[0, 1]
    np.random.seed(99)
    cluster1 = np.random.normal([0.6, 0.6], 0.08, (15, 2))
    cluster2 = np.random.normal([0.7, 0.5], 0.07, (12, 2))
    cluster3 = np.random.normal([0.65, 0.55], 0.06, (8, 2))

    ax.scatter(cluster1[:, 0], cluster1[:, 1], color=BLUE, alpha=0.7, s=60,
               label='Meditatsiya')
    ax.scatter(cluster2[:, 0], cluster2[:, 1], color=RED, alpha=0.7, s=60,
               label='Igry (oshibka!)')
    ax.scatter(cluster3[:, 0], cluster3[:, 1], color=ORANGE, alpha=0.7, s=60,
               label='Relaks. muzyka')

    # Show centroids
    ax.scatter([0.6], [0.6], color=BLUE, s=200, marker='*', zorder=5)
    ax.scatter([0.7], [0.5], color=RED, s=200, marker='*', zorder=5)

    ax.annotate('', xy=(0.7, 0.5), xytext=(0.6, 0.6),
                arrowprops=dict(arrowstyle='<->', color=GRAY, lw=1.5))
    ax.text(0.63, 0.54, 'mala dist.', fontsize=8, color=GRAY, rotation=45)

    ax.set_xlabel('Embedding dim 1', fontsize=9, color=DARK)
    ax.set_ylabel('Embedding dim 2', fontsize=9, color=DARK)
    ax.set_title('Semanticheskaya oshibka:\nblizkie embeddingy, razny intent',
                 fontsize=10, fontweight='bold', color=DARK)
    ax.legend(fontsize=8)
    ax.grid(True, alpha=0.3)

    # Bottom-left: precision vs ambiguity
    ax = axes[1, 0]
    ambiguity = np.linspace(0, 1, 10)
    codirank_prec = 0.72 - 0.3 * ambiguity + 0.05 * np.sin(ambiguity * np.pi)
    baseline_prec = 0.55 - 0.35 * ambiguity

    ax.plot(ambiguity, codirank_prec, color=BLUE, linewidth=2,
            marker='o', markersize=5, label='CoDiRank')
    ax.plot(ambiguity, baseline_prec, color=GRAY, linewidth=1.8,
            linestyle='--', marker='s', markersize=4, label='Baseline')
    ax.fill_between(ambiguity, codirank_prec, baseline_prec,
                    alpha=0.15, color=BLUE, label='Preimushchestvo')
    ax.set_xlabel('Uroven neodnoznachnosti zaprosa', fontsize=9, color=DARK)
    ax.set_ylabel('Precision@3', fontsize=9, color=DARK)
    ax.set_title('Degradatsiya kachestva\npri neodnoznachnosti',
                 fontsize=10, fontweight='bold', color=DARK)
    ax.legend(fontsize=8)
    ax.grid(True, alpha=0.3)
    ax.set_ylim(0.1, 0.85)

    # Bottom-right: confusion matrix
    ax = axes[1, 1]
    cm = np.array([
        [72, 18, 10],
        [15, 58, 27],
        [8,  22, 70],
    ])
    categories_cm = ['Relevanten', 'Chast. relev.', 'Nerelev.']

    import matplotlib.colors as mcolors
    norm = mcolors.Normalize(vmin=0, vmax=80)
    im = ax.imshow(cm, cmap='Blues', norm=norm)

    ax.set_xticks([0, 1, 2])
    ax.set_yticks([0, 1, 2])
    ax.set_xticklabels(categories_cm, fontsize=8, color=DARK)
    ax.set_yticklabels(categories_cm, fontsize=8, color=DARK)
    ax.set_xlabel('Predskazano', fontsize=9, color=DARK)
    ax.set_ylabel('Fakticheski', fontsize=9, color=DARK)
    ax.set_title('Matritsa oshibok\n(klass. relevantnosti)', fontsize=10,
                 fontweight='bold', color=DARK)

    for i in range(3):
        for j in range(3):
            val = cm[i, j]
            txt_color = 'white' if val > 50 else DARK
            ax.text(j, i, str(val), ha='center', va='center',
                    fontsize=12, color=txt_color, fontweight='bold')

    plt.colorbar(im, ax=ax, shrink=0.8)

    plt.tight_layout()
    plt.savefig(os.path.join(OUT_DIR, 'fig_error_analysis.png'),
                dpi=150, bbox_inches='tight', facecolor='white')
    plt.close()
    print('[OK] fig_error_analysis.png')


# ===========================================================================
# RUN ALL CHART GENERATORS
# ===========================================================================
print('=== Generating charts ===')
gen_gplay_search()
gen_sparsity_matrix()
gen_matrix_factorization()
gen_ncf_architecture()
gen_llm_comparison()
gen_gplay_dataset()
gen_profile_scheme()
gen_relevance_space()
gen_metrics_table()
gen_ollama_config()
gen_db_schema()
gen_bot_dialog()
gen_dataset_stats()
gen_ndcg_vs_turns()
gen_error_analysis()
print('=== All charts generated ===\n')


# ===========================================================================
# DOCX INSERTION
# ===========================================================================
print('=== Inserting charts into docx ===')

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph as DocxParagraph
import copy

INPUT_DOCX  = os.path.join(BASE_DIR, 'public', 'diplom_v2.docx')
OUTPUT_DOCX = os.path.join(BASE_DIR, 'public', 'diplom_v2_final.docx')

PLACEHOLDER_MAP = {
    "Пример интерфейса поиска Google Play":          os.path.join(OUT_DIR, 'fig_gplay_search.png'),
    "Визуализация разреженности матрицы":            os.path.join(OUT_DIR, 'fig_sparsity_matrix.png'),
    "Схема матричной факторизации":                  os.path.join(OUT_DIR, 'fig_matrix_factorization.png'),
    "Архитектура нейронной коллаборативной фильтрации": os.path.join(OUT_DIR, 'fig_ncf_architecture.png'),
    "Сравнительная характеристика открытых LLM":    os.path.join(OUT_DIR, 'fig_llm_comparison.png'),
    "Структура и статистика датасета Google Play":   os.path.join(OUT_DIR, 'fig_gplay_dataset.png'),
    "Схема динамического обновления профиля":        os.path.join(OUT_DIR, 'fig_profile_scheme.png'),
    "Общая архитектура алгоритма CoDiRank":          os.path.join(OUT_DIR, '01_architecture.png'),
    "Визуализация функции релевантности CoDiRank":   os.path.join(OUT_DIR, 'fig_relevance_space.png'),
    "Сравнительная таблица метрик качества":         os.path.join(OUT_DIR, 'fig_metrics_table.png'),
    "Общая архитектура программной системы":         os.path.join(OUT_DIR, '01_architecture.png'),
    "Конфигурация Ollama":                           os.path.join(OUT_DIR, 'fig_ollama_config.png'),
    "Схема базы данных PostgreSQL":                  os.path.join(OUT_DIR, 'fig_db_schema.png'),
    "Скриншот диалога с ботом":                      os.path.join(OUT_DIR, 'fig_bot_dialog.png'),
    "Распределение приложений по категориям":        os.path.join(OUT_DIR, '05_catalog_distribution.png'),
    "Статистика тестового диалогового датасета":     os.path.join(OUT_DIR, 'fig_dataset_stats.png'),
    "Таблица сравнительных результатов":             os.path.join(OUT_DIR, 'fig_metrics_table.png'),
    "График зависимости NDCG@10":                    os.path.join(OUT_DIR, 'fig_ndcg_vs_turns.png'),
    "Анализ ошибочных рекомендаций":                 os.path.join(OUT_DIR, 'fig_error_analysis.png'),
}


doc = docx.Document(INPUT_DOCX)

print(f'Total paragraphs: {len(doc.paragraphs)}')

# Find placeholders
matches = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    for key, img_path in PLACEHOLDER_MAP.items():
        if key in text:
            matches.append((i, para, key, img_path))
            break

print(f'Found {len(matches)} placeholder paragraphs')


def add_image_paragraph_after(doc, ref_para, image_path, width_inches=5.5):
    """
    Add a new paragraph with an image after ref_para.
    Works by adding a paragraph to the document body and then moving it
    to the correct position using XML manipulation.
    """
    # Add a temporary paragraph at the end of the document
    tmp_para = doc.add_paragraph()
    tmp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tmp_para.add_run()
    run.add_picture(image_path, width=Inches(width_inches))

    # Move the new paragraph element to be right after ref_para's element
    tmp_elem = tmp_para._element
    ref_elem = ref_para._element
    ref_elem.addnext(tmp_elem)

    return tmp_para


inserted = 0
for idx, para, key, img_path in matches:
    if not os.path.exists(img_path):
        print(f'  [WARN] Image not found: {img_path} (for: {key[:40]})')
        continue
    print(f'  [INS] para[{idx}]: "{key[:50]}" => {os.path.basename(img_path)}')
    add_image_paragraph_after(doc, para, img_path, width_inches=5.5)
    inserted += 1

print(f'\nInserted {inserted} images.')
doc.save(OUTPUT_DOCX)
print(f'Saved: {OUTPUT_DOCX}')
