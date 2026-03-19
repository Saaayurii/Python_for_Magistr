#!/usr/bin/env python3
"""Insert remaining charts into the diploma document at appropriate sections."""
from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path(__file__).parent.parent
PUBLIC = BASE / "public"
INPUT_DOCX  = PUBLIC / "diplom_v2_final.docx"
OUTPUT_DOCX = PUBLIC / "diplom_v2_final.docx"

# (anchor_fragment, image_file, caption, insert_after_n_paragraphs)
# anchor_fragment: unique text to search for the insertion point paragraph
# insert_after: how many paragraphs after the anchor to place the image
INSERTIONS = [
    # 02_fsm_dialog → after FSM description in 3.1.4
    (
        "конечного автомата состояний (FSM) для управления диал",
        "02_fsm_dialog.png",
        "Рисунок — Граф состояний конечного автомата (FSM) диалога бота",
        2,
    ),
    # 03_profile_update → after the EMA formula paragraph (2.1.2)
    (
        "P(t) = alpha * P(t-1) + (1 - alpha) * E(at) * w(at)",
        "03_profile_update.png",
        "Рисунок — Динамика нормы вектора профиля P(t) и влияние весов тональности",
        1,
    ),
    # 04_ranking_formula → after R(i,D,t) definition paragraph (2.2.3)
    (
        "Функция релевантности R(i,D,t) для элемента каталога i на шаге диалога t",
        "04_ranking_formula.png",
        "Рисунок — Декомпозиция оценки R(i,D,t) и веса компонентов ранжирования",
        2,
    ),
    # 06_cold_start_comparison → after cold start experiment paragraph (3.3.4)
    (
        "Для оценки эффективности обработки проблемы холодного старта проведён",
        "06_cold_start_comparison.png",
        "Рисунок — Сравнение методов рекомендаций в условиях холодного старта",
        2,
    ),
    # 07_embeddings_2d → after LLM/embedding description paragraph (1.3.2)
    (
        "Применение больших языковых моделей для извлечения предпочтений",
        "07_embeddings_2d.png",
        "Рисунок — Визуализация семантического пространства эмбеддингов (t-SNE) и траектория профиля P(t)",
        3,
    ),
    # 08_dialog_metrics → after experiment result paragraph (3.3.3)
    (
        "3.3.3 Анализ влияния глубины диалога на точность профиля",
        "08_dialog_metrics.png",
        "Рисунок — Распределение длины диалогов до рекомендации и структура обратной связи",
        3,
    ),
    # 09_alpha_sensitivity → after alpha description in profile model (2.1.2)
    (
        "Профиль интересов пользователя P(t) на шаге диалога t представляется",
        "09_alpha_sensitivity.png",
        "Рисунок — Влияние коэффициента инерции α на скорость сходимости профиля пользователя",
        3,
    ),
    # 10_tech_stack → after system components description (3.1 intro)
    (
        "Программная реализация системы построена на принципах модульности",
        "10_tech_stack.png",
        "Рисунок — Технологический стек системы CoDiRank",
        2,
    ),
]


def find_paragraph_index(doc: docx.Document, fragment: str) -> int | None:
    for i, para in enumerate(doc.paragraphs):
        if fragment.lower() in para.text.lower():
            return i
    return None


def insert_image_after(doc: docx.Document, para_idx: int, image_path: Path,
                       caption: str, width: float = 5.5) -> None:
    """Insert centered image + caption after paragraph at para_idx."""
    anchor_para = doc.paragraphs[para_idx]

    # Add paragraphs at end of doc (so they inherit the document part)
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(str(image_path), width=Inches(width))

    cap_para = doc.add_paragraph()
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_para.add_run(caption)
    cap_run.italic = True
    cap_run.font.size = Pt(10)

    # Move both XML elements to the correct position (after anchor)
    # cap first (so it ends up after img after two addnext calls)
    anchor_para._element.addnext(cap_para._element)
    anchor_para._element.addnext(img_para._element)


def main() -> None:
    doc = docx.Document(str(INPUT_DOCX))
    inserted = 0

    for anchor_text, img_file, caption, skip in INSERTIONS:
        img_path = PUBLIC / img_file
        if not img_path.exists():
            print(f"  SKIP (not found): {img_file}")
            continue

        idx = find_paragraph_index(doc, anchor_text)
        if idx is None:
            print(f"  SKIP (anchor not found): {anchor_text[:60]}...")
            continue

        # reload paragraphs after each insertion (indices shift)
        insert_at = idx + skip
        if insert_at >= len(doc.paragraphs):
            insert_at = idx

        insert_image_after(doc, insert_at, img_path, caption)
        print(f"  OK [{insert_at}] {img_file}: {caption[:60]}")
        inserted += 1

    doc.save(str(OUTPUT_DOCX))
    print(f"\nDone: {inserted} charts inserted → {OUTPUT_DOCX.name}")


if __name__ == "__main__":
    print(f"Loading {INPUT_DOCX.name}...\n")
    main()
